#!/usr/bin/env python3
"""Meeting One-Pager — Current Project Status (Jan 31, 2026)"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# Title
title = doc.add_heading('Vesla ERP — Project Status', level=1)
title.runs[0].font.color.rgb = RGBColor(47, 84, 150)
title.runs[0].font.size = Pt(16)

sub = doc.add_paragraph('January 31, 2026  |  Phase 1 Review')
sub.runs[0].font.size = Pt(9)
sub.runs[0].font.color.rgb = RGBColor(128, 128, 128)
sub.runs[0].font.italic = True

# --- System Health ---
h = doc.add_heading('System Health', level=2)
h.runs[0].font.size = Pt(12)
items = [
    ('Backend (port 3000):', ' ✅ Healthy — DB connected, all core packages OK, uptime stable'),
    ('Staff Admin App (port 8083):', ' ✅ Live and serving'),
    ('Frontend (port 5173):', ' ✅ Running (React + Vite)'),
    ('Database (Neon):', ' ✅ Connected — 275 Prisma models, 628 indexes'),
]
for label, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label); run.bold = True; run.font.size = Pt(10)
    r2 = p.add_run(desc); r2.font.size = Pt(10)

# --- Recent Deliveries ---
h = doc.add_heading('Shipped This Week (Jan 25–31)', level=2)
h.runs[0].font.size = Pt(12)
shipped = [
    'TARS Phase 1 API completion — full endpoint coverage',
    'Customer Support Package — tickets, reviews, mobile support, seeded test data',
    'Tech Support page — interactive ticket creation from admin panel',
    'Sidebar restructure — Rental Admin merged into Rent-A-Car, cleaner navigation',
    'Marketing module — WhatsApp Business Campaigns integration',
    'General Ledger, Trial Balance, Profit & Loss, Balance Sheet pages',
    'Contract Renewals — moved to Rent-A-Car section, permissions seeded',
    '85 sidebar permissions seeded across all modules',
    'Vehicle Clarity wired to real backend data',
    'Brute Force Testing Platform — stress test all ERP modules',
    'Support ticket number generator fixed (global max, not per-company)',
]
for item in shipped:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(item)
    run.font.size = Pt(10)

# --- By The Numbers ---
h = doc.add_heading('By The Numbers', level=2)
h.runs[0].font.size = Pt(12)
numbers = [
    ('651', 'commits this week (since Jan 25)'),
    ('94', 'kanban tasks completed'),
    ('23', 'items in review'),
    ('41', 'items in backlog'),
    ('525', 'commits by Anders (backend, infra, APIs)'),
    ('106', 'active branches'),
]
for num, desc in numbers:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(num); run.bold = True; run.font.size = Pt(10)
    r2 = p.add_run(f' — {desc}'); r2.font.size = Pt(10)

# --- In Review / In Progress ---
h = doc.add_heading('Key Items In Review', level=2)
h.runs[0].font.size = Pt(12)
review_items = [
    ('Anders:', 'TARS staging verification, Dispatch Map, stub services implementation, customer support ticketing, E2E test suite, superuser logic, permissions to business logic, CI & branch protection'),
    ('Mia:', 'Contract Renewal Engine, MRM landing page, Rent-A-Car Mobile review, Blog deployment, UAE Pass VHD signing, CCAvenue payment hardcoded IDs (P0), permission gating on routes (P0), Kevin service account'),
    ('Miguel:', 'UAE Pass Email (in-progress)'),
]
for person, items_str in review_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(person); run.bold = True; run.font.size = Pt(10)
    r2 = p.add_run(f' {items_str}'); r2.font.size = Pt(10)

# --- Open Risks ---
h = doc.add_heading('Open Risks & Blockers', level=2)
h.runs[0].font.size = Pt(12)
risks = [
    ('P0 Security:', ' No permission gating on Admin/Finance/HR routes — needs fix before any client deployment'),
    ('P0 Security:', ' Health/monitoring endpoints exposed without auth'),
    ('P0 Finance:', ' Hardcoded payment account IDs in CCAvenue routes'),
    ('Kevin Support:', ' No service account yet — cannot monitor support tickets via API'),
    ('Backend Stubs:', ' 7 core services still have throw-stub logic (booking, rental contract, invoice, payment, etc.)'),
]
for label, desc in risks:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label); run.bold = True; run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(192, 0, 0)
    r2 = p.add_run(desc); r2.font.size = Pt(10)

# --- Cloud Cost Summary ---
h = doc.add_heading('Cloud Cost Projection (5 Companies, Medium Load)', level=2)
h.runs[0].font.size = Pt(12)
costs = [
    ('Starter ($81/mo):', ' 1 CU autoscale + Render Standard — good for onboarding'),
    ('Growth ($244/mo):', ' 2 CU + Render Pro — for 5 active daily clients'),
    ('Enterprise ($671/mo):', ' Always-on + replica + load balanced — SLA/compliance'),
    ('Per-client breakdown:', ' Available in Vesla_ERP_Cloud_Cost_Projection_v2.xlsx'),
]
for label, desc in costs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label); run.bold = True; run.font.size = Pt(10)
    r2 = p.add_run(desc); r2.font.size = Pt(10)

# Footer
doc.add_paragraph()
footer = doc.add_paragraph('Prepared by Kevin 🔧  |  January 31, 2026  |  Data pulled live from codebase + kanban')
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

out = "/Users/miguelitodeguzman/Projects/tech-project/docs/Vesla_ERP_Meeting_Status_2026-01-31.docx"
doc.save(out)
print(f"✅ Saved: {out}")
