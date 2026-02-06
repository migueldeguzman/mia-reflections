#!/usr/bin/env python3
"""Generate professional Word documents from Phase 1 markdown files."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# Colors
HEADER_BG = "2E74B5"
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW_BG = "F2F2F2"
DARK_BLUE = RGBColor(0x2E, 0x74, 0xB5)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=10, font_name="Calibri", align=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    # Set spacing
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def set_cell_rich_text(cell, text, bold=False, color=None, size=10, font_name="Calibri"):
    """Set cell text with bold segments parsed from **text**."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    # Parse bold segments
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('~~') and part.endswith('~~'):
            run = p.add_run(part[2:-2])
            run.font.strike = True
        else:
            # Handle strikethrough within non-bold text
            sub_parts = re.split(r'(~~.*?~~)', part)
            for sub in sub_parts:
                if sub.startswith('~~') and sub.endswith('~~'):
                    run = p.add_run(sub[2:-2])
                    run.font.strike = True
                else:
                    run = p.add_run(sub)
        run.font.name = font_name
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, header, bold=True, color=HEADER_TEXT, size=10)
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)
            set_cell_rich_text(cell, cell_text, size=10)
    
    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)
    
    # Set all cell vertical alignment to top and remove excessive padding
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
            tcPr.append(tcVAlign)
    
    return table


def add_heading(doc, text, level=1):
    """Add a heading with Calibri font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_BLUE if level <= 2 else RGBColor(0x33, 0x33, 0x33)


def add_para(doc, text, bold=False, italic=False, size=11):
    """Add a paragraph with rich text formatting (parses **bold**)."""
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
            if bold:
                run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        if italic:
            run.italic = True
    return p


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * level)
    return p


def set_default_font(doc):
    """Set default document font to Calibri."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ============================================================
# DOCUMENT 1: PHASE_1_REMAINING.docx
# ============================================================

def generate_remaining():
    doc = Document()
    set_default_font(doc)
    
    # Title
    title = doc.add_heading('Phase 1 — What Changed & What Remains', level=0)
    for run in title.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_BLUE
    
    add_para(doc, '**Updated: January 31, 2026 — 7:35 AM (Presentation Day)**')
    add_para(doc, '**Compiled by Kevin 🔧**')
    
    # Separator
    doc.add_paragraph('─' * 60)
    
    # === Overnight Changes ===
    add_heading(doc, 'Overnight Changes (Jan 30 evening → Jan 31 morning)', level=1)
    
    # TS Error Reduction
    add_heading(doc, '🔥 Massive TypeScript Error Reduction (Anders overnight)', level=2)
    add_styled_table(doc,
        ['Metric', 'Yesterday', 'Today', 'Change'],
        [
            ['Frontend TS errors', '468', '52', '-416 errors killed ✅'],
            ['Backend TS errors', '0', '0', 'Clean ✅'],
        ],
        col_widths=[5, 3, 3, 6]
    )
    
    add_para(doc, 'Anders pushed **15 commits** overnight in 9 batches, crushing ~500+ frontend TypeScript errors. The codebase went from 468 errors to just 52 remaining.')
    
    # New Changes Overnight
    add_heading(doc, '✅ New Changes Overnight', level=2)
    add_styled_table(doc,
        ['Change', 'Detail'],
        [
            ['Support tickets module permissions', 'New permission set for support ticket access control'],
            ['Vehicle TARS service', '**Dropped from 17 throws → 4 throws** — Anders implemented it (was full stub)'],
            ['MFA routes registered', "Previously existed but weren't mounted — now live with frontend permission guards"],
            ['Sidebar navigation', 'Recovery and Fleet Partners modules added (were missing)'],
            ['Customer support package', 'Item tagging, public feedback URLs, image uploads'],
            ['Support issue categories', 'Admin-configurable with 10 defaults'],
            ['ReportIssueWidget', 'Now POSTs to mrm-server correctly'],
            ['Duplicate Forgot Password', 'Fixed (was showing twice on login page)'],
        ],
        col_widths=[6, 11]
    )
    
    # Kanban Board Movement
    add_heading(doc, 'Kanban Board Movement', level=2)
    add_styled_table(doc,
        ['Status', 'Yesterday', 'Today', 'Change'],
        [
            ['Done', '87', '87', '—'],
            ['Backlog', '42', '43', '+1'],
            ['Review', '18', '**24**', '**+6**'],
            ['In Progress', '4', '3', '-1'],
            ['Total', '172', '**179**', '+7 new cards'],
        ],
        col_widths=[4, 3, 3, 4]
    )
    
    # Separator
    doc.add_paragraph('─' * 60)
    
    # === What Still Remains ===
    add_heading(doc, 'What STILL Remains', level=1)
    
    # Backend Service Stubs
    add_heading(doc, 'Backend Service Stubs (Throw Counts)', level=2)
    add_styled_table(doc,
        ['Service', 'Yesterday', 'Today', 'Change'],
        [
            ['user.service.ts', '34', '**34**', '—'],
            ['role.service.ts', '26', '**26**', '—'],
            ['superuser.service.ts', '24', '**24**', '—'],
            ['invoice.service.ts', '22', '**22**', '—'],
            ['booking.service.ts', '15', '**15**', '—'],
            ['transaction.service.ts', '15', '**15**', '—'],
            ['recovery.service.ts', '11', '**11**', '—'],
            ['vehicle-dealership.service.ts', '6', '**6**', '—'],
            ['vehicle-tars.service.ts', '17', '**4**', '**-13** ✅'],
            ['properties.service.ts', '4', '**4**', '—'],
        ],
        col_widths=[6, 3, 3, 4]
    )
    
    # Infrastructure
    add_heading(doc, 'Infrastructure', level=2)
    add_styled_table(doc,
        ['Item', 'Status'],
        [
            ['CI Pipeline', '**Still all failing** — latest 5 runs all `failure`'],
            ['Frontend TS errors', '**52 remaining** (was 468) — massive improvement ✅'],
            ['Backend TS errors', '**0** ✅'],
            ['Backend server (mrm-server)', 'Miguel & Mia working on deployment'],
            ['Branch protection', 'Not enforced'],
        ],
        col_widths=[6, 11]
    )
    
    # Modules Still Early Stage
    add_heading(doc, 'Modules Still Early Stage', level=2)
    add_bullet(doc, '**Properties** — dashboard only, service stub (4 throws)')
    add_bullet(doc, '**Vehicle Dealership** — dashboard only, service stub (6 throws)')
    add_bullet(doc, '**Showroom Mobile** — empty repo (0 files)')
    add_bullet(doc, '**Vendor Provider App** — empty repo (0 files)')
    
    # Separator
    doc.add_paragraph('─' * 60)
    
    # === Summary: Yesterday vs Today ===
    add_heading(doc, 'Summary: Yesterday vs Today', level=1)
    add_styled_table(doc,
        ['Metric', 'Yesterday (5:30 PM)', 'Today (7:35 AM)', 'Verdict'],
        [
            ['P0 Security Issues', '0 (fixed yesterday)', '0', '✅ Clean'],
            ['Frontend TS Errors', '468', '**52**', '**🔥 -89%**'],
            ['Backend TS Errors', '0', '0', '✅'],
            ['Vehicle TARS Stubs', '17', '**4**', '✅ Mostly implemented'],
            ['MFA', 'Routes not registered', '**Registered + gated**', '✅'],
            ['Support Module', 'Basic', '**Permissions + categories + tagging + uploads**', '✅'],
            ['Sidebar Nav', 'Missing Recovery/Fleet', '**All modules visible**', '✅'],
            ['Kanban Cards', '172', '179', '+7 new'],
            ['Remaining Service Stubs', '~45+', '~40+', 'Slight improvement'],
        ],
        col_widths=[5, 5, 5, 4]
    )
    
    # Separator
    doc.add_paragraph('─' * 60)
    
    # === Presentation Readiness ===
    add_heading(doc, 'Presentation Readiness', level=1)
    
    add_heading(doc, 'Stronger than yesterday:', level=2)
    add_bullet(doc, 'Frontend codebase 89% cleaner (52 vs 468 errors)')
    add_bullet(doc, 'Backend at zero TS errors')
    add_bullet(doc, 'All modules visible in sidebar navigation')
    add_bullet(doc, 'Support system significantly enhanced')
    add_bullet(doc, 'TARS service mostly implemented')
    add_bullet(doc, 'MFA functional and gated')
    
    add_heading(doc, 'Still avoid in live demo:', level=2)
    add_bullet(doc, 'End-to-end booking → invoice → payment flow (core services still stubbed)')
    add_bullet(doc, "TARS fines page if data still won't load")
    add_bullet(doc, 'Package Management table')
    add_bullet(doc, 'CI metrics')
    
    add_heading(doc, 'Recommended approach:', level=2)
    add_para(doc, 'Guided UI walkthrough + architecture overview + mobile app demos + "what\'s ahead" roadmap from the Gaps & Opportunities doc.')
    
    # Separator
    doc.add_paragraph('─' * 60)
    
    # Closing
    p = add_para(doc, 'Anders crushed it overnight. 15 commits, 416 TypeScript errors eliminated. The codebase is in the best shape it\'s been.', italic=True)
    
    # Save
    output = '/Users/miguelitodeguzman/Projects/tech-project/docs/PHASE_1_REMAINING.docx'
    doc.save(output)
    print(f"✅ Saved: {output}")


# ============================================================
# DOCUMENT 2: PHASE_1_GAPS_AND_OPPORTUNITIES.docx
# ============================================================

def generate_gaps():
    doc = Document()
    set_default_font(doc)
    
    # Title
    title = doc.add_heading('Vesla ERP — Gap Analysis & Opportunities', level=0)
    for run in title.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_BLUE
    
    add_para(doc, '**What We Have, What We\'re Missing, What We Can Add**')
    add_para(doc, '**Prepared for presentation — January 31, 2026 | Kevin 🔧**')
    
    doc.add_paragraph('─' * 60)
    
    # Purpose
    add_heading(doc, 'Purpose', level=1)
    add_para(doc, 'This document maps Vesla ERP against **UAE market requirements** and **industry-standard features** for rent-a-car, fleet management, and ERP businesses. Organized into three categories:')
    add_bullet(doc, '✅ **We Have It** — built and in Phase 1')
    add_bullet(doc, '⚠️ **Started But Incomplete** — code exists, needs completion')
    add_bullet(doc, '🆕 **Opportunity To Add** — not yet built, would differentiate us')
    
    doc.add_paragraph('─' * 60)
    
    # === Section 1: UAE Regulatory ===
    add_heading(doc, '1. UAE REGULATORY & COMPLIANCE', level=1)
    add_styled_table(doc,
        ['Feature', 'Status', 'Detail'],
        [
            ['VAT (5%) — Tax invoices, VAT returns, tax categories', '✅ Have It', 'VAT Returns page, Tax Categories, tax codes, tax rates all in Finance module'],
            ['Corporate Tax (9% above AED 375K) — Tax period management', '✅ Have It', 'Corporate Tax page + ct_tax_periods model, tax loss tracking'],
            ['E-Invoicing — FTA-compliant electronic invoicing', '⚠️ Started', 'Routes exist (export + transmission, 377 lines), credentials model. UAE FTA rollout ongoing — needs completion for compliance'],
            ['WPS (Wage Protection System) — Salary file generation', '⚠️ Started', 'Prisma models: wps_submissions, wps_agents, wps_errors. Payroll routes exist. Full WPS file format compliance TBD'],
            ['RTA Integration — Traffic fines & Salik charges', '⚠️ Started', 'TARS module built (UI + routes + models). Auto-sync from RTA not connected. Data loading currently fails'],
            ['UAE Pass — Government digital identity', '⚠️ Started', 'VHD PDF Signing routes + sessions exist. Full UAE Pass SSO integration TBD'],
            ['Economic Substance Regulations (ESR)', '🆕 Opportunity', 'UAE requires certain entities to demonstrate economic substance. Reporting module would serve holding companies'],
            ['AML/KYC Compliance — Anti-money laundering', '⚠️ Started', 'KYC documents module exists (routes, UI). Full AML screening/reporting TBD'],
            ['MOHRE Integration — Labour contracts & WPS', '🆕 Opportunity', 'Direct integration with Ministry of Human Resources for visa/labour card tracking'],
        ],
        col_widths=[5, 3, 9]
    )
    
    # === Section 2: Rent-a-Car Operations ===
    add_heading(doc, '2. RENT-A-CAR OPERATIONS', level=1)
    add_styled_table(doc,
        ['Feature', 'Status', 'Detail'],
        [
            ['Booking Management', '⚠️ Started', 'Full UI built (dashboard, chat, approvals). Backend service has 15 stubs — core flow incomplete'],
            ['Contract Management', '⚠️ Started', 'UI complete (create, extend, upgrade, void). Backend service has 12 stubs'],
            ['Vehicle Fleet CRUD', '✅ Have It', 'Full vehicle management (images, setup, config, blocks, sync)'],
            ['Delivery & Dispatch', '✅ Have It', 'Slots, calendar, tracking, driver efficiency. Mobile dispatch in progress'],
            ['Dynamic Pricing / Yield Management', '✅ Have It', '9 pages, pricing rules, calculator, history, variables. One of our strongest modules'],
            ['Customer-Facing Booking Portal', '✅ Have It', 'Rent-a-Car Mobile app (137 tsx files, 88 screens) — customer-facing booking experience'],
            ['Online Reservation Widget', '🆕 Opportunity', "Embeddable booking widget for the client's website. Most competitors offer this"],
            ['Aggregator Integration (Kayak, Skyscanner, Rentalcars.com)', '🆕 Opportunity', 'No integration with travel aggregator platforms. Major revenue channel for rental companies'],
            ['Digital Contract E-Signing', '⚠️ Started', 'Signature service exists (10 stubs). Not functional yet'],
            ['Customer Self-Service App', '✅ Have It', 'Rent-a-Car Mobile (137 tsx, 88 screens). Showroom Mobile repo exists for future expansion'],
            ['Multi-Location / Branch Management', '⚠️ Partial', 'companyId isolation exists (308 service files). No dedicated branch/location management with pickup/dropoff points'],
            ['Replacement Vehicle Workflow', '✅ Have It', 'Replace vehicle/driver modals in rental operations'],
            ['Contract Renewal Engine', '⚠️ In Review', 'Mia has this in review — renewal offers, settings, dashboard page exists'],
        ],
        col_widths=[5, 3, 9]
    )
    
    # === Section 3: Fleet & Vehicle Management ===
    add_heading(doc, '3. FLEET & VEHICLE MANAGEMENT', level=1)
    add_styled_table(doc,
        ['Feature', 'Status', 'Detail'],
        [
            ['Vehicle Lifecycle Tracking', '⚠️ Started', 'Routes + models exist. Analytics service unclear'],
            ['Vehicle Insurance Management', '⚠️ Started', 'Controller, routes, service exist — but service has 8 stubs'],
            ['GPS/Telematics Integration', '⚠️ Started', 'Tracking routes, dispatch tracking components, location controllers exist. No telematics provider integration (Wialon, Geotab, etc.)'],
            ['Fuel Management / Fuel Cards', '🆕 Opportunity', 'References in codebase but no dedicated fuel tracking module. Fleet companies need fuel card integration (ENOC, ADNOC)'],
            ['Tire Management', '🆕 Opportunity', 'References exist but no dedicated tire lifecycle tracking. Important for fleet cost control'],
            ['Preventive Maintenance Scheduling', '⚠️ Started', 'Maintenance schedule service exists (7 stubs). Service center module has work orders'],
            ['Vehicle Inspection Checklists', '⚠️ Started', 'Inspection items in service center models. No mobile check-in/check-out inspection flow'],
            ['Vehicle Purchase & Procurement', '⚠️ Started', 'Purchase order (10 stubs) and purchase request (6 stubs) services exist'],
            ['TCO (Total Cost of Ownership)', '🆕 Opportunity', 'No total cost tracking per vehicle (purchase + maintenance + fuel + insurance + depreciation). Valuable for fleet decision-making'],
            ['Fleet Utilization Analytics', '⚠️ Started', 'Vehicle Clarity now wired to real data (fixed today). Fleet analytics service has 4 stubs'],
            ['Driver Behavior Monitoring', '🆕 Opportunity', 'Would require telematics integration. Harsh braking, speeding, idle time'],
            ['Vehicle Remarketing / Disposal', '🆕 Opportunity', 'End-of-lifecycle vehicle sales/auction management. Connects to Dealership module'],
        ],
        col_widths=[5, 3, 9]
    )
    
    # === Section 4: Finance & Accounting ===
    add_heading(doc, '4. FINANCE & ACCOUNTING', level=1)
    add_styled_table(doc,
        ['Feature', 'Status', 'Detail'],
        [
            ['Invoicing', '⚠️ Started', 'Full UI, 22 backend stubs'],
            ['Accounts Receivable', '✅ Have It', 'Vouchers, aging report, receivables management'],
            ['Accounts Payable', '⚠️ Started', 'Service has 8 stubs'],
            ['Chart of Accounts', '✅ Have It', 'Full CoA management'],
            ['Fixed Assets & Depreciation', '✅ Have It', 'Asset tracking + depreciation calculations'],
            ['Collection & Legal Cases', '✅ Have It', 'Queue, notices, legal case management'],
            ['Bank Integration / Auto-Reconciliation', '⚠️ Started', 'Bank accounts & reconciliation UI exists. Auto-import from banks not built'],
            ['Multi-Currency', '⚠️ Partial', 'Currency references throughout (217 files). Dedicated multi-currency transaction handling TBD'],
            ['Budget Management', '🆕 Opportunity', 'No budget module. Companies need departmental budgets, actuals vs budget tracking'],
            ['Cash Flow Forecasting', '🆕 Opportunity', 'Forecast references exist (7 files) but no dedicated cash flow projection tool'],
            ['Financial Dashboards / BI', '⚠️ Started', 'Finance dashboard exists. No advanced BI/analytics with drill-down'],
            ['Audit Trail', '✅ Have It', 'Audit logs in admin module'],
            ['Payment Gateway', '⚠️ Started', 'CCAvenue integrated (test mode). P0 hardcoded IDs being fixed'],
            ['Revenue Recognition', '🆕 Opportunity', 'For rental contracts spanning multiple periods — IFRS 16 compliance'],
        ],
        col_widths=[5, 3, 9]
    )
    
    # === Section 5: HR & Workforce ===
    add_heading(doc, '5. HR & WORKFORCE', level=1)
    add_styled_table(doc,
        ['Feature', 'Status', 'Detail'],
        [
            ['Attendance Tracking', '✅ Have It', 'Attendance page in HR module'],
            ['Leave Management', '✅ Have It', 'Leave management page'],
            ['Payroll Processing', '⚠️ Started', 'Routes + models exist. WPS compliance depth unclear'],
            ['Gratuity / End of Service', '🆕 Opportunity', 'UAE law requires end-of-service gratuity calculation. Not built yet'],
            ['Employee Self-Service Portal', '🆕 Opportunity', "Employees can't view payslips, request leave, or update info through a portal"],
            ['Document Management (visa, labor card, Emirates ID)', '🆕 Opportunity', 'Critical for UAE businesses — visa expiry tracking, labor card renewal alerts'],
            ['Biometric / Time Clock Integration', '🆕 Opportunity', 'Most UAE companies use biometric attendance. Integration would automate time tracking'],
            ['Training & Certification Tracking', '🆕 Opportunity', 'Driver training records, safety certifications'],
            ['Accommodation Management', '🆕 Opportunity', 'Many UAE companies provide staff housing — tracking capacity, assignments'],
            ['Performance Management', '🆕 Opportunity', 'KPIs, reviews, promotion tracking'],
        ],
        col_widths=[5, 3, 9]
    )
    
    # === Section 6: Customer Experience ===
    add_heading(doc, '6. CUSTOMER EXPERIENCE & COMMUNICATION', level=1)
    add_styled_table(doc,
        ['Feature', 'Status', 'Detail'],
        [
            ['Support Ticket System', '✅ Have It', 'Dashboard, detail page, comments, attachments, SLA configs, escalation rules'],
            ['Knowledge Base / FAQ', '✅ Have It', '197 entries across 21 modules'],
            ['WhatsApp Business Integration', '⚠️ Started', 'Routes, templates, settings exist. Full send/receive integration TBD'],
            ['SMS Notifications', '⚠️ Partial', 'Notification infrastructure exists (94 files). SMS gateway integration TBD'],
            ['Email Notifications', '⚠️ Started', 'Email verification works. Transactional emails (booking confirmation, invoice, reminder) TBD'],
            ['Customer CRM', '🆕 Opportunity', 'No CRM module. Customer history, preferences, communication log, segmentation'],
            ['Loyalty / Rewards Program', '🆕 Opportunity', 'Only 2 references in codebase. Points, tiers, rewards would drive retention'],
            ['Customer Mobile App', '🆕 Opportunity', 'Showroom Mobile repo is empty. Customer-facing app for booking + payments + docs'],
            ['Rating & Review System', '🆕 Opportunity', 'Post-rental feedback collection. Drives service quality'],
            ['AI Chatbot / Automated Support', '🆕 Opportunity', 'FAQ exists but no automated chat. AI-powered first response would reduce support load'],
            ['Multi-Channel Communication Hub', '🆕 Opportunity', 'Unified inbox (WhatsApp + SMS + email + in-app) for customer communication'],
        ],
        col_widths=[5, 3, 9]
    )
    
    # === Section 7: Technology & Platform ===
    add_heading(doc, '7. TECHNOLOGY & PLATFORM', level=1)
    add_styled_table(doc,
        ['Feature', 'Status', 'Detail'],
        [
            ['Multi-Tenant SaaS', '✅ Have It', 'companyId isolation, subscription system, package access middleware'],
            ['Role-Based Access Control', '✅ Have It', '299 permissions, granular module-level (compartmentalized today)'],
            ['Mobile Apps', '⚠️ Partial', 'Rent-a-Car (137 tsx), Service Center (50 tsx). Recovery/Kanban early. Showroom/Vendor empty'],
            ['CI/CD Pipeline', '⚠️ Broken', '6 workflows exist but all failing. Branch protection not enforced'],
            ['API Documentation', '✅ Have It', 'Swagger at /api-docs'],
            ['Public API / API Marketplace', '🆕 Opportunity', 'No external API for third-party integrations. Would enable partner ecosystem'],
            ['Offline Mode for Field Agents', '🆕 Opportunity', 'Recovery drivers, delivery agents, service technicians need offline capability'],
            ['Report Builder / Custom Reports', '🆕 Opportunity', 'Fixed report pages exist. No drag-and-drop custom report builder'],
            ['Data Export (Excel/PDF)', '⚠️ Partial', 'Some export exists. Comprehensive export across all modules TBD'],
            ['Notification Center', '⚠️ Started', 'Push token models exist. Unified in-app notification center TBD'],
            ['Dark Mode / Theme Customization', '✅ Have It', 'Theme editor in admin'],
            ['Audit & Activity Logging', '✅ Have It', 'Audit logs module'],
        ],
        col_widths=[5, 3, 9]
    )
    
    doc.add_paragraph('─' * 60)
    
    # === Top 10 Opportunities ===
    add_heading(doc, 'Top 10 Opportunities for Future Phases', level=1)
    add_para(doc, 'These are the biggest gaps that would differentiate Vesla in the UAE market:')
    
    add_styled_table(doc,
        ['#', 'Opportunity', 'Business Impact', 'Effort'],
        [
            ['1', '~~Customer-Facing Booking Portal~~', '✅ Already built — Rent-a-Car Mobile (137 tsx, 88 screens)', 'Done'],
            ['2', '**Gratuity & End-of-Service Calculator**', 'UAE legal requirement — every company needs this', 'Low'],
            ['3', '**GPS/Telematics Integration** (Wialon, Geotab)', 'Real-time fleet visibility, stolen vehicle recovery', 'Medium'],
            ['4', '~~Customer Mobile App~~', '✅ Already built — Rent-a-Car Mobile', 'Done'],
            ['5', '**Fuel Card Integration** (ENOC/ADNOC)', 'Automated fuel cost tracking per vehicle', 'Medium'],
            ['6', '**Employee Self-Service Portal**', 'Reduce HR admin burden, payslips, leave requests', 'Medium'],
            ['7', '**Aggregator Integration** (Kayak/Rentalcars.com)', 'Massive distribution channel for bookings', 'Medium'],
            ['8', '**E-Invoicing Completion** (UAE FTA)', 'Regulatory — will become mandatory', 'Medium'],
            ['9', '**Multi-Channel Communication Hub**', 'Unified WhatsApp + SMS + email inbox', 'Medium'],
            ['10', '**TCO / Fleet Analytics Dashboard**', 'Data-driven fleet decisions (buy/sell/maintain)', 'Medium'],
        ],
        col_widths=[1, 5, 7, 2.5]
    )
    
    doc.add_paragraph('─' * 60)
    
    # === Summary for Presentation ===
    add_heading(doc, 'Summary for Presentation', level=1)
    
    # Add blockquote-style summary
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    quote_text = (
        '"Vesla ERP covers 13 packages across the full business lifecycle of a UAE rent-a-car and fleet management company. '
        'Phase 1 delivers strong frontend coverage across all modules, with backend depth in Finance, Admin, Service Center, and Dynamic Pricing. '
        'The platform is multi-tenant, permission-gated, and UAE-compliant (VAT, Corporate Tax, TARS). '
        'The roadmap includes customer-facing portals, telematics integration, and deeper UAE regulatory compliance — '
        'positioning Vesla as the only UAE-built, UAE-focused ERP for the rent-a-car industry."'
    )
    run = p.add_run(quote_text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    
    doc.add_paragraph('─' * 60)
    
    p = add_para(doc, 'Based on codebase analysis + UAE regulatory research + industry competitor benchmarking (RENTALL, Fleetio, industry standards). Compiled by Kevin 🔧 on Day 1 — Mia & Anders may have additional context.', italic=True, size=10)
    
    # Save
    output = '/Users/miguelitodeguzman/Projects/tech-project/docs/PHASE_1_GAPS_AND_OPPORTUNITIES.docx'
    doc.save(output)
    print(f"✅ Saved: {output}")


if __name__ == '__main__':
    generate_remaining()
    generate_gaps()
    print("\n🎉 Both documents generated successfully!")
