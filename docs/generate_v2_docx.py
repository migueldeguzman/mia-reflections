#!/usr/bin/env python3
"""
Generate .docx versions of PHASE_1_PRESENTATION_v2.md and CLOUD_COST_PROJECTION_v2.md
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_from_md(doc, header_row, data_rows, style='Table Grid'):
    """Add a formatted table to the document."""
    if not header_row:
        return
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols, style=style)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = cell_text.strip()
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2B5797')

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < cols:
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = cell_text.strip()
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        if '**' in cell_text:
                            run.bold = True
                            run.text = run.text.replace('**', '')
                # Alternate row shading
                if row_idx % 2 == 0:
                    set_cell_shading(cell, 'E8EDF3')

    doc.add_paragraph()  # spacing after table


def parse_md_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (header, rows, end_idx)."""
    header = []
    rows = []
    idx = start_idx

    # Parse header
    if idx < len(lines) and '|' in lines[idx]:
        parts = [p.strip() for p in lines[idx].split('|')]
        header = [p for p in parts if p and not all(c in '-| ' for c in p)]
        idx += 1

    # Skip separator
    if idx < len(lines) and '|' in lines[idx] and '---' in lines[idx]:
        idx += 1

    # Parse rows
    while idx < len(lines) and '|' in lines[idx]:
        parts = [p.strip() for p in lines[idx].split('|')]
        row = [p for p in parts if p != '']
        # Handle empty cells
        raw = lines[idx].strip()
        if raw.startswith('|'):
            raw = raw[1:]
        if raw.endswith('|'):
            raw = raw[:-1]
        row = [c.strip() for c in raw.split('|')]
        if row and not all(c == '' or all(ch in '-| ' for ch in c) for c in row):
            rows.append(row)
        idx += 1

    return header, rows, idx


def clean_text(text):
    """Remove markdown formatting from text."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = text.replace('✅', '[DONE]').replace('🟡', '[WIP]').replace('🔲', '[PLANNED]').replace('🔴', '[CRITICAL]').replace('🟢', '[NICE]')
    # Keep emojis that work in docx
    return text.strip()


def add_formatted_paragraph(doc, text, style='Normal'):
    """Add a paragraph with basic bold/italic support."""
    text = clean_text(text)
    p = doc.add_paragraph(style=style)
    # Simple: just add the text
    run = p.add_run(text)
    return p


def md_to_docx(md_path, docx_path, title):
    """Convert a markdown file to a formatted .docx document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    idx = 0
    in_code_block = False
    code_lines = []

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            idx += 1
            continue

        # Code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                p.paragraph_format.left_indent = Cm(1)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            idx += 1
            continue

        if in_code_block:
            code_lines.append(line)
            idx += 1
            continue

        # Tables
        if '|' in stripped and not stripped.startswith('>'):
            # Check if next line is separator
            if idx + 1 < len(lines) and '---' in lines[idx + 1]:
                header, rows, new_idx = parse_md_table(lines, idx)
                if header:
                    add_table_from_md(doc, header, rows)
                idx = new_idx
                continue

        # Headings
        if stripped.startswith('# '):
            text = clean_text(stripped[2:])
            doc.add_heading(text, level=0)
            idx += 1
            continue
        elif stripped.startswith('## '):
            text = clean_text(stripped[3:])
            doc.add_heading(text, level=1)
            idx += 1
            continue
        elif stripped.startswith('### '):
            text = clean_text(stripped[4:])
            doc.add_heading(text, level=2)
            idx += 1
            continue
        elif stripped.startswith('#### '):
            text = clean_text(stripped[5:])
            doc.add_heading(text, level=3)
            idx += 1
            continue

        # Horizontal rules
        if stripped == '---' or stripped == '***':
            # Add a thin line / page break indicator
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            idx += 1
            continue

        # Blockquotes
        if stripped.startswith('>'):
            text = clean_text(stripped[1:].strip())
            p = doc.add_paragraph(style='Intense Quote') if 'Intense Quote' in [s.name for s in doc.styles] else doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x2B, 0x57, 0x97)
            idx += 1
            continue

        # List items
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = clean_text(stripped[2:])
            p = doc.add_paragraph(text, style='List Bullet')
            idx += 1
            continue
        elif re.match(r'^\d+\.\s', stripped):
            text = clean_text(re.sub(r'^\d+\.\s', '', stripped))
            p = doc.add_paragraph(text, style='List Number')
            idx += 1
            continue

        # Regular paragraph (skip metadata lines like *Last updated*)
        if stripped.startswith('*') and stripped.endswith('*') and len(stripped) > 2:
            text = stripped.strip('*').strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            idx += 1
            continue

        # Default: regular paragraph
        text = clean_text(stripped)
        if text:
            doc.add_paragraph(text)

        idx += 1

    doc.save(docx_path)
    print(f"  ✓ Generated: {docx_path}")


def main():
    print("Generating .docx files from v2 markdown documents...\n")

    # Document 1: Phase 1 Presentation v2
    md1 = os.path.join(DOCS_DIR, 'PHASE_1_PRESENTATION_v2.md')
    docx1 = os.path.join(DOCS_DIR, 'PHASE_1_PRESENTATION_v2.docx')
    md_to_docx(md1, docx1, 'Vesla ERP — Phase 1 Status Presentation')

    # Document 2: Cloud Cost Projection v2
    md2 = os.path.join(DOCS_DIR, 'CLOUD_COST_PROJECTION_v2.md')
    docx2 = os.path.join(DOCS_DIR, 'CLOUD_COST_PROJECTION_v2.docx')
    md_to_docx(md2, docx2, 'Vesla ERP — Cloud Cost Projection v2')

    print(f"\nDone! Files saved to: {DOCS_DIR}")


if __name__ == '__main__':
    main()
