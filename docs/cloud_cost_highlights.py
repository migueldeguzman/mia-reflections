#!/usr/bin/env python3
"""Generate Cloud Cost Highlights one-pager as .docx"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Vesla ERP — Cloud Cost Highlights', level=1)
title.runs[0].font.color.rgb = RGBColor(47, 84, 150)

sub = doc.add_paragraph('5 Companies  |  Medium Load  |  Neon + Render')
sub.runs[0].font.size = Pt(10)
sub.runs[0].font.color.rgb = RGBColor(128, 128, 128)
sub.runs[0].font.italic = True

# --- What We're Hosting ---
doc.add_heading("What We're Hosting", level=2)
items = [
    ('Backend:', 'Node.js + Express + Prisma (275 models, 628 indexes)'),
    ('Frontend:', 'React + Vite (static site)'),
    ('Database:', 'PostgreSQL on Neon (serverless)'),
    ('Workers:', 'TARS sync, Speed Sync, cron jobs'),
]
for label, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label + ' ')
    run.bold = True
    p.add_run(desc)

# --- Per Company Profile ---
doc.add_heading('Per Company Profile (Medium Load)', level=2)
profiles = [
    '100–300 vehicles',
    '10–25 concurrent staff users',
    '300–800 contracts/month',
    '~3–5 GB database after Year 1',
]
for item in profiles:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Total for 5 companies: ~15–25 GB database')
run.bold = True

# --- Monthly Cost Options ---
doc.add_heading('Monthly Cost Options', level=2)

# Option A
p = doc.add_heading('Option A: Starter — $81/mo  ($16/company)', level=3)
option_a = [
    'Neon Launch: $49  (1 CU autoscale, scale-to-zero nights/weekends)',
    'Render Standard backend: $25  (1 CPU, 2 GB RAM)',
    'Render Starter worker: $7',
    'Frontend: free  (static site)',
]
for item in option_a:
    doc.add_paragraph(item, style='List Bullet')

# Option B
p = doc.add_heading('Option B: Growth — $244/mo  ($49/company)', level=3)
option_b = [
    'Neon Launch: $115  (2 CU autoscale, heavier traffic)',
    'Render Pro backend: $85  (2 CPU, 4 GB RAM)',
    'Render Standard worker: $25',
    'Team workspace: $19',
]
for item in option_b:
    doc.add_paragraph(item, style='List Bullet')

# Option C
p = doc.add_heading('Option C: Enterprise — $671/mo  ($134/company)', level=3)
option_c = [
    'Neon Scale: $419  (always-on + read replica, SOC2/SLA)',
    'Render Pro ×2 backend: $170  (load balanced)',
    'Render Standard worker: $25',
    'Team workspace ×3: $57',
]
for item in option_c:
    doc.add_paragraph(item, style='List Bullet')

# --- Key Takeaways ---
doc.add_heading('Key Takeaways', level=2)
takeaways = [
    ('Start at $81/mo', ' — enough for 5 medium clients onboarding'),
    ('Biggest cost driver:', ' Neon compute (usage-based, quiet hours = free)'),
    ('Storage is cheap:', ' $0.35/GB — even 25 GB is only $8.75/mo'),
    ('Scale trigger:', ' upgrade when CPU consistently >70% or response times degrade'),
    ('Not included:', ' domains, email service, file storage (S3), third-party APIs (RTA, payment gateways)'),
]
for label, desc in takeaways:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label)
    run.bold = True
    p.add_run(desc)

# Footer
doc.add_paragraph()
footer = doc.add_paragraph('Prepared by Kevin 🔧  |  January 31, 2026')
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

out = "/Users/miguelitodeguzman/Projects/tech-project/docs/Vesla_ERP_Cloud_Cost_Highlights.docx"
doc.save(out)
print(f"✅ Saved to {out}")
