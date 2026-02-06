#!/usr/bin/env python3
"""
Generate Vesla ERP Deployment Readiness Checklist (.docx)
Professional document with checkboxes, color-coded priorities, and file references.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ── Color constants ──
P0_COLOR = RGBColor(0xCC, 0x00, 0x00)  # Red
P1_COLOR = RGBColor(0xE6, 0x7E, 0x00)  # Orange
P2_COLOR = RGBColor(0x00, 0x70, 0xC0)  # Blue
HEADER_BG = "1F4E79"
SUBHEADER_BG = "D6E4F0"
P0_BG = "FDE0DC"
P1_BG = "FFF3E0"
P2_BG = "E3F2FD"
DONE_BG = "E8F5E9"
FILE_COLOR = RGBColor(0x4A, 0x4A, 0x4A)
CHECK = "☐"


def set_cell_bg(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_cell_text(cell, text, bold=False, color=None, size=9, alignment=None):
    """Add formatted text to a cell."""
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def priority_tag(priority):
    """Return color info for priority."""
    if priority == "P0":
        return P0_COLOR, P0_BG
    elif priority == "P1":
        return P1_COLOR, P1_BG
    else:
        return P2_COLOR, P2_BG


def add_section_header(doc, title, number):
    """Add a major section header."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def add_subsection(doc, title):
    """Add a subsection heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def create_checklist_table(doc, items):
    """
    Create a styled checklist table.
    items: list of dicts with keys: check, priority, item, owner, file_ref, notes
    """
    col_widths = [Cm(1.2), Cm(1.3), Cm(8.0), Cm(2.2), Cm(4.5), Cm(1.5)]
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    headers = ["Status", "Priority", "Item", "Owner", "File / Reference", "Notes"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, HEADER_BG)
        add_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for item in items:
        row = table.add_row()
        pri = item.get("priority", "P2")
        pri_color, pri_bg = priority_tag(pri)

        # Status (checkbox)
        c0 = row.cells[0]
        add_cell_text(c0, CHECK, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Priority
        c1 = row.cells[1]
        set_cell_bg(c1, pri_bg)
        add_cell_text(c1, pri, bold=True, color=pri_color, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Item description
        c2 = row.cells[2]
        p = c2.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item["item"])
        run.font.size = Pt(9)

        # Owner
        c3 = row.cells[3]
        add_cell_text(c3, item.get("owner", "—"), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # File reference
        c4 = row.cells[4]
        ref_text = item.get("file_ref", "—")
        add_cell_text(c4, ref_text, size=7, color=FILE_COLOR)

        # Notes
        c5 = row.cells[5]
        add_cell_text(c5, item.get("notes", ""), size=7)

    # Set column widths
    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = width

    return table


def build_document():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(29.7)  # A4 landscape
    section.page_height = Cm(21.0)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── Styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # ═══════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("VESLA ERP")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Deployment Readiness Checklist")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("UAE Rent-a-Car & Fleet Management ERP Platform")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_paragraph()

    # Meta table
    meta_table = doc.add_table(rows=7, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Document Version", "1.0"),
        ("Date", datetime.datetime.now().strftime("%B %d, %Y")),
        ("Classification", "CONFIDENTIAL"),
        ("Stack", "Node.js + Express + Prisma + React + Vite + Expo"),
        ("Database", "Neon PostgreSQL (275 models, 628 indexes)"),
        ("Deploy Target", "Render (staging: develop, production: main)"),
        ("Repository", "web-erp-app"),
    ]
    for i, (k, v) in enumerate(meta_data):
        c0 = meta_table.rows[i].cells[0]
        c1 = meta_table.rows[i].cells[1]
        set_cell_bg(c0, SUBHEADER_BG)
        add_cell_text(c0, k, bold=True, size=10)
        add_cell_text(c1, v, size=10)
        c0.width = Cm(5)
        c1.width = Cm(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # TABLE OF CONTENTS / LEGEND
    # ═══════════════════════════════════════════════════════════
    toc_title = doc.add_paragraph()
    run = toc_title.add_run("How to Use This Checklist")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()

    # Priority legend
    legend = doc.add_table(rows=4, cols=3)
    legend.alignment = WD_TABLE_ALIGNMENT.LEFT
    legend_headers = ["Priority", "Meaning", "Action Required"]
    for i, h in enumerate(legend_headers):
        set_cell_bg(legend.rows[0].cells[i], HEADER_BG)
        add_cell_text(legend.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)

    legend_data = [
        ("P0", "BLOCKER — Must fix before production launch", "Blocks deployment. No exceptions."),
        ("P1", "HIGH — Should fix before production, critical for operations", "Fix in Sprint 0 or deploy with documented risk acceptance."),
        ("P2", "MEDIUM — Should fix within 30 days of launch", "Can launch with these open; schedule immediately post-launch."),
    ]
    for i, (pri, meaning, action) in enumerate(legend_data):
        row = legend.rows[i + 1]
        pri_color, pri_bg = priority_tag(pri)
        set_cell_bg(row.cells[0], pri_bg)
        add_cell_text(row.cells[0], pri, bold=True, color=pri_color, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(row.cells[1], meaning, size=10)
        add_cell_text(row.cells[2], action, size=10)
        row.cells[0].width = Cm(2)
        row.cells[1].width = Cm(10)
        row.cells[2].width = Cm(8)

    doc.add_paragraph()

    status_p = doc.add_paragraph()
    run = status_p.add_run("Status Legend:  ")
    run.bold = True
    run.font.size = Pt(10)
    run = status_p.add_run("☐ = Not Started    ☑ = Complete    ⚠ = In Progress / Blocked")
    run.font.size = Pt(10)

    doc.add_paragraph()

    sections_list = doc.add_paragraph()
    run = sections_list.add_run("Sections:")
    run.bold = True
    run.font.size = Pt(12)
    sections = [
        "1. Infrastructure & Platform",
        "2. Security (P0 Priority)",
        "3. Backend Services",
        "4. Frontend Application",
        "5. Mobile Application",
        "6. Testing & Quality",
        "7. UAE Compliance & Regulatory",
        "8. Monitoring & Operations",
        "9. Go-Live & Launch",
        "10. Summary & Sign-Off",
    ]
    for s in sections:
        p = doc.add_paragraph(s, style='List Bullet')
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 1: INFRASTRUCTURE
    # ═══════════════════════════════════════════════════════════
    add_section_header(doc, "INFRASTRUCTURE & PLATFORM", "1")

    add_subsection(doc, "1.1  Database Readiness (Neon PostgreSQL)")
    create_checklist_table(doc, [
        {"priority": "P0", "item": "All Prisma migrations applied to production DB (prisma migrate deploy)", "owner": "Backend", "file_ref": "backend/prisma/migrations/", "notes": ""},
        {"priority": "P0", "item": "Migration state matches schema.prisma — no drift", "owner": "Backend", "file_ref": "backend/prisma/schema.prisma", "notes": ""},
        {"priority": "P0", "item": "Run seed:production (core + permissions + business + integrations)", "owner": "Backend", "file_ref": "npm run seed:production", "notes": ""},
        {"priority": "P1", "item": "Verify all 628 indexes exist and are not redundant", "owner": "DBA", "file_ref": "schema.prisma @@index directives", "notes": ""},
        {"priority": "P1", "item": "Connection pooling configured (Neon serverless driver + Prisma)", "owner": "Backend", "file_ref": "backend/src/lib/prisma.ts", "notes": ""},
        {"priority": "P1", "item": "DATABASE_URL uses ?sslmode=require", "owner": "DevOps", "file_ref": "render.yaml → envVars", "notes": ""},
        {"priority": "P2", "item": "Neon branching tested for staging isolation", "owner": "DevOps", "file_ref": "npm run neon:branch:create", "notes": ""},
        {"priority": "P2", "item": "Database backup schedule configured (daily snapshots)", "owner": "DevOps", "file_ref": "npm run snapshot:daily", "notes": ""},
        {"priority": "P2", "item": "Backup restore tested on staging", "owner": "DevOps", "file_ref": "npm run snapshot:restore", "notes": ""},
        {"priority": "P1", "item": "Tenant isolation verified (multi-company companyId filters)", "owner": "Backend", "file_ref": "npm run db:verify-isolation", "notes": ""},
    ])

    doc.add_paragraph()

    add_subsection(doc, "1.2  Render Configuration")
    create_checklist_table(doc, [
        {"priority": "P0", "item": "render.yaml production service configured and valid", "owner": "DevOps", "file_ref": "render.yaml", "notes": ""},
        {"priority": "P0", "item": "Build command verified: install → prisma generate → migrate deploy → seed:default → build frontend → build backend", "owner": "DevOps", "file_ref": "render.yaml → buildCommand", "notes": ""},
        {"priority": "P0", "item": "Start command: npm start --workspace=backend", "owner": "DevOps", "file_ref": "render.yaml → startCommand", "notes": ""},
        {"priority": "P0", "item": "All required env vars set in Render Dashboard (DATABASE_URL, JWT_SECRET, JWT_REFRESH_SECRET, CORS_ORIGIN, BACKUP_ENCRYPTION_KEY)", "owner": "DevOps", "file_ref": "render.yaml → envVars (sync: false)", "notes": ""},
        {"priority": "P0", "item": "DISABLE_RATE_LIMIT = false in production", "owner": "DevOps", "file_ref": "render.yaml line 18", "notes": ""},
        {"priority": "P0", "item": "ENABLE_SUPERUSER_BYPASS = false in production", "owner": "DevOps", "file_ref": "render.yaml line 20", "notes": ""},
        {"priority": "P0", "item": "NODE_ENV = production", "owner": "DevOps", "file_ref": "render.yaml line 14", "notes": ""},
        {"priority": "P1", "item": "Staging service deploys from develop branch", "owner": "DevOps", "file_ref": "render.yaml → vesla-erp-staging", "notes": ""},
        {"priority": "P1", "item": "Auto-deploy enabled for staging, manual for production", "owner": "DevOps", "file_ref": "Render Dashboard", "notes": ""},
        {"priority": "P2", "item": "Health check path configured in Render", "owner": "DevOps", "file_ref": "/api/health", "notes": ""},
    ])

    doc.add_paragraph()

    add_subsection(doc, "1.3  Domain, DNS & SSL")
    create_checklist_table(doc, [
        {"priority": "P1", "item": "Custom domain configured (e.g., app.vesla.ae)", "owner": "DevOps", "file_ref": "Render Dashboard → Custom Domains", "notes": ""},
        {"priority": "P1", "item": "DNS A/CNAME records point to Render", "owner": "DevOps", "file_ref": "DNS provider", "notes": ""},
        {"priority": "P1", "item": "SSL/TLS certificate provisioned (Render auto-SSL)", "owner": "DevOps", "file_ref": "Render Dashboard", "notes": ""},
        {"priority": "P1", "item": "HTTPS redirect enforced (no plain HTTP)", "owner": "DevOps", "file_ref": "Render settings", "notes": ""},
        {"priority": "P2", "item": "CDN configured for /assets/* (cache-control headers set in render.yaml)", "owner": "DevOps", "file_ref": "render.yaml → headers", "notes": "max-age=31536000"},
        {"priority": "P2", "item": "API domain matches CORS_ORIGIN configuration", "owner": "Backend", "file_ref": "backend/src/services/cors-config.service.ts", "notes": ""},
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 2: SECURITY
    # ═══════════════════════════════════════════════════════════
    add_section_header(doc, "SECURITY (P0 PRIORITY)", "2")

    p = doc.add_paragraph()
    run = p.add_run("⚠ CRITICAL: These items were flagged in the security audit. No production launch without all P0 items resolved.")
    run.bold = True
    run.font.color.rgb = P0_COLOR
    run.font.size = Pt(10)

    doc.add_paragraph()

    add_subsection(doc, "2.1  Authentication & Authorization")
    create_checklist_table(doc, [
        {"priority": "P0", "item": "Permission gating on ALL Admin routes (/admin/*)", "owner": "Backend", "file_ref": "backend/src/routes/admin.routes.ts + admin/", "notes": "KNOWN GAP"},
        {"priority": "P0", "item": "Permission gating on ALL Finance routes (/finance/*)", "owner": "Backend", "file_ref": "backend/src/routes/accounting.routes.ts", "notes": "KNOWN GAP"},
        {"priority": "P0", "item": "Permission gating on ALL HR routes (/hr/*)", "owner": "Backend", "file_ref": "backend/src/routes/ (HR routes)", "notes": "KNOWN GAP"},
        {"priority": "P0", "item": "JWT_SECRET is unique, ≥32 chars, not default", "owner": "DevOps", "file_ref": "Render env → JWT_SECRET", "notes": ""},
        {"priority": "P0", "item": "JWT_REFRESH_SECRET is different from JWT_SECRET", "owner": "DevOps", "file_ref": "Render env → JWT_REFRESH_SECRET", "notes": ""},
        {"priority": "P0", "item": "Token expiry: 1h access / 7d refresh in production", "owner": "Backend", "file_ref": "backend/src/services/auth.service.ts", "notes": ""},
        {"priority": "P1", "item": "CSRF middleware active on state-changing endpoints", "owner": "Backend", "file_ref": "backend/src/middleware/csrf.middleware.ts", "notes": ""},
        {"priority": "P1", "item": "Cross-company IDOR fix verified (P0-1 pentest passed)", "owner": "Security", "file_ref": "P0_SECURITY_FIXES_PENTEST_REPORT.md", "notes": "✅ Fixed"},
        {"priority": "P1", "item": "UUID validation middleware on all param routes", "owner": "Backend", "file_ref": "backend/src/middleware/auth.middleware.ts", "notes": "✅ Fixed"},
    ])

    doc.add_paragraph()

    add_subsection(doc, "2.2  Secrets & Configuration")
    create_checklist_table(doc, [
        {"priority": "P0", "item": "Remove hardcoded CCAvenue merchant_id=50496 from test file", "owner": "Backend", "file_ref": "backend/src/services/payment/__test-ccavenue.ts", "notes": "KNOWN ISSUE"},
        {"priority": "P0", "item": "All CCAvenue credentials loaded from env vars / DB config", "owner": "Backend", "file_ref": "backend/src/services/payment/payment.service.ts", "notes": "Verify gateway config"},
        {"priority": "P0", "item": "No secrets in source code (grep audit passed)", "owner": "Security", "file_ref": "Run: grep -r 'password\\|secret\\|api_key' src/", "notes": ""},
        {"priority": "P0", "item": "Speed Scraper credentials in env vars only", "owner": "DevOps", "file_ref": "render.yaml → SPEED_SCRAPER_EMAIL/PASSWORD", "notes": ""},
        {"priority": "P1", "item": ".env files in .gitignore", "owner": "Backend", "file_ref": ".gitignore", "notes": ""},
        {"priority": "P1", "item": "BACKUP_ENCRYPTION_KEY set (not empty)", "owner": "DevOps", "file_ref": "Render env", "notes": ""},
        {"priority": "P2", "item": "API key rotation plan documented", "owner": "Security", "file_ref": "Create: docs/api-key-rotation.md", "notes": ""},
    ])

    doc.add_paragraph()

    add_subsection(doc, "2.3  Rate Limiting, CORS & Endpoints")
    create_checklist_table(doc, [
        {"priority": "P0", "item": "Rate limiting enabled in production (DISABLE_RATE_LIMIT=false)", "owner": "DevOps", "file_ref": "render.yaml", "notes": ""},
        {"priority": "P0", "item": "Health endpoints require auth or are scoped (not exposing internals)", "owner": "Backend", "file_ref": "/api/health/detailed, /api/health/database", "notes": "KNOWN GAP"},
        {"priority": "P0", "item": "CORS_ORIGIN is specific domain (not wildcard *)", "owner": "Backend", "file_ref": "backend/src/services/cors-config.service.ts", "notes": ""},
        {"priority": "P1", "item": "Rate limiter configured per-route (auth: stricter, read: relaxed)", "owner": "Backend", "file_ref": "backend/src/middleware/rate-limiter.middleware.ts", "notes": ""},
        {"priority": "P1", "item": "Sync endpoints rate limited", "owner": "Backend", "file_ref": "backend/src/middleware/sync-rate-limiter.middleware.ts", "notes": ""},
        {"priority": "P1", "item": "File upload size limits enforced", "owner": "Backend", "file_ref": "backend/src/middleware/upload.middleware.ts", "notes": ""},
        {"priority": "P1", "item": "Input validation middleware on all POST/PUT routes", "owner": "Backend", "file_ref": "backend/src/middleware/validation.middleware.ts", "notes": ""},
        {"priority": "P1", "item": "Response sanitizer prevents data leakage", "owner": "Backend", "file_ref": "backend/src/middleware/response-sanitizer.middleware.ts", "notes": ""},
        {"priority": "P2", "item": "Swagger UI disabled in production (unless ENABLE_SWAGGER=true)", "owner": "Backend", "file_ref": "DEPLOYMENT.md → Environment Comparison", "notes": ""},
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 3: BACKEND
    # ═══════════════════════════════════════════════════════════
    add_section_header(doc, "BACKEND SERVICES", "3")

    add_subsection(doc, "3.1  Stub Services (Must Implement or Defer)")

    p = doc.add_paragraph()
    run = p.add_run("The following backend services are partially or fully stubbed. Each must be implemented, marked as Phase 2, or explicitly deferred with risk acceptance.")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    create_checklist_table(doc, [
        {"priority": "P0", "item": "rental-contract-document.service.ts — FULLY STUBBED (upload/download/list return errors)", "owner": "Backend", "file_ref": "backend/src/services/rental-contract-document.service.ts", "notes": "Schema models missing"},
        {"priority": "P1", "item": "driver-profile.service.ts — Stubbed until schema updated", "owner": "Backend", "file_ref": "backend/src/services/driver-profile.service.ts", "notes": "Schema fields needed"},
        {"priority": "P1", "item": "vehicle-grn.service.ts — Fixed Asset & Accounts Payable creation STUBBED", "owner": "Backend", "file_ref": "backend/src/services/vehicle-grn.service.ts:283-287", "notes": "Models not in schema"},
        {"priority": "P1", "item": "vehicle-lifecycle-tracker.service.ts — Some models stubbed to null", "owner": "Backend", "file_ref": "backend/src/services/vehicle-lifecycle-tracker.service.ts:122", "notes": ""},
        {"priority": "P1", "item": "superuser.service.ts — mrmi_audit_logs model STUBBED (3 locations)", "owner": "Backend", "file_ref": "backend/src/services/superuser.service.ts:87,983,1002", "notes": "Model not in schema"},
        {"priority": "P2", "item": "customer-documents.service.ts — Contains stub references", "owner": "Backend", "file_ref": "backend/src/services/customer-documents.service.ts", "notes": ""},
        {"priority": "P2", "item": "notification-dispatch.service.ts — Stub references", "owner": "Backend", "file_ref": "backend/src/services/notification-dispatch.service.ts", "notes": ""},
        {"priority": "P2", "item": "password-reset.service.ts — Verify email integration is live", "owner": "Backend", "file_ref": "backend/src/services/password-reset.service.ts", "notes": "SMTP required"},
        {"priority": "P2", "item": "platform-settings.service.ts — Stub references", "owner": "Backend", "file_ref": "backend/src/services/platform-settings.service.ts", "notes": ""},
        {"priority": "P2", "item": "vehicle-lifecycle-analytics.service.ts — Stub references", "owner": "Backend", "file_ref": "backend/src/services/vehicle-lifecycle-analytics.service.ts", "notes": ""},
    ])

    doc.add_paragraph()

    add_subsection(doc, "3.2  Database & Seed Data")
    create_checklist_table(doc, [
        {"priority": "P0", "item": "seed:default runs successfully on clean DB", "owner": "Backend", "file_ref": "backend/prisma/seeds/seed-default.ts", "notes": ""},
        {"priority": "P0", "item": "Superuser account created with secure password", "owner": "Backend", "file_ref": "npm run admin:setup-master", "notes": ""},
        {"priority": "P0", "item": "All permission packages seeded (seed:permissions)", "owner": "Backend", "file_ref": "backend/prisma/seeds/02-permissions/", "notes": ""},
        {"priority": "P1", "item": "Chart of accounts seeded", "owner": "Backend", "file_ref": "backend/prisma/seeds/08-chart-of-accounts/", "notes": ""},
        {"priority": "P1", "item": "Payment gateway config seeded (CCAvenue)", "owner": "Backend", "file_ref": "backend/prisma/seeds/06-integrations/seed-ccavenue.ts", "notes": ""},
        {"priority": "P1", "item": "FTA compliance workflows seeded", "owner": "Backend", "file_ref": "npm run seed:fta-workflows", "notes": ""},
        {"priority": "P1", "item": "Seed idempotency verified (can run multiple times)", "owner": "Backend", "file_ref": "npm run db:verify-seeds", "notes": ""},
    ])

    doc.add_paragraph()

    add_subsection(doc, "3.3  Error Handling, Logging & Background Jobs")
    create_checklist_table(doc, [
        {"priority": "P1", "item": "Error middleware catches all unhandled errors", "owner": "Backend", "file_ref": "backend/src/middleware/error.middleware.ts", "notes": ""},
        {"priority": "P1", "item": "Structured logging in production (JSON format, no console.log)", "owner": "Backend", "file_ref": "DEPLOYMENT.md → Console drops in prod", "notes": ""},
        {"priority": "P1", "item": "Audit logging persists to database", "owner": "Backend", "file_ref": "backend/src/services/audit.service.ts", "notes": "P1-3 fix verified"},
        {"priority": "P1", "item": "TARS sync job runs on schedule", "owner": "Backend", "file_ref": "tars/ directory", "notes": "Fines & Salik sync"},
        {"priority": "P1", "item": "Speed Sync scraper configured and tested", "owner": "Backend", "file_ref": "speed-scraper/speed_scraper.py", "notes": "Python + Selenium"},
        {"priority": "P1", "item": "Background job failure alerts configured", "owner": "DevOps", "file_ref": "backend/src/services/tars-sync-log.service.ts", "notes": ""},
        {"priority": "P2", "item": "Request context middleware provides trace IDs", "owner": "Backend", "file_ref": "backend/src/middleware/request-context.middleware.ts", "notes": ""},
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 4: FRONTEND
    # ═══════════════════════════════════════════════════════════
    add_section_header(doc, "FRONTEND APPLICATION", "4")

    add_subsection(doc, "4.1  Mock Data Pages (Must Wire to Backend)")

    p = doc.add_paragraph()
    run = p.add_run("The following 16 frontend pages use hardcoded mock/sample data instead of API calls. They must be wired to real backend endpoints before launch or explicitly marked as Phase 2.")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    create_checklist_table(doc, [
        # Finance pages
        {"priority": "P0", "item": "FinanceActivityPage.tsx — mockActivities array (269 lines of mock data)", "owner": "Frontend", "file_ref": "frontend/src/pages/finance/FinanceActivityPage.tsx", "notes": ""},
        {"priority": "P0", "item": "CollectionNoticesPage.tsx — mockNotices array", "owner": "Frontend", "file_ref": "frontend/src/pages/finance/CollectionNoticesPage.tsx", "notes": ""},
        {"priority": "P1", "item": "ARVouchersPage.tsx — mock AR voucher data", "owner": "Frontend", "file_ref": "frontend/src/pages/finance/receivables/ARVouchersPage.tsx", "notes": ""},
        {"priority": "P1", "item": "ARVoucherDetailPage.tsx — mock AR voucher detail", "owner": "Frontend", "file_ref": "frontend/src/pages/finance/receivables/ARVoucherDetailPage.tsx", "notes": ""},
        {"priority": "P1", "item": "ARAgingReportPage.tsx — mock aging report data", "owner": "Frontend", "file_ref": "frontend/src/pages/finance/reports/ARAgingReportPage.tsx", "notes": ""},
        {"priority": "P1", "item": "BankReconciliationPage.tsx — mock bank data", "owner": "Frontend", "file_ref": "frontend/src/pages/finance/treasury/BankReconciliationPage.tsx", "notes": ""},
        {"priority": "P1", "item": "DepositsPage.tsx — mock deposits data", "owner": "Frontend", "file_ref": "frontend/src/pages/finance/treasury/DepositsPage.tsx", "notes": ""},
        {"priority": "P1", "item": "BankAccountsPage.tsx — mock/stub bank accounts", "owner": "Frontend", "file_ref": "frontend/src/pages/finance/treasury/BankAccountsPage.tsx", "notes": ""},
        # Collection pages
        {"priority": "P1", "item": "CollectionQueuePage.tsx — mock queue data", "owner": "Frontend", "file_ref": "frontend/src/pages/finance/collection/CollectionQueuePage.tsx", "notes": ""},
        {"priority": "P1", "item": "InvoiceMatching.tsx — mock matching data", "owner": "Frontend", "file_ref": "frontend/src/pages/finance/collection/InvoiceMatching.tsx", "notes": ""},
        {"priority": "P1", "item": "CollectionEntryDetailPage.tsx — mock detail data", "owner": "Frontend", "file_ref": "frontend/src/pages/finance/collection/CollectionEntryDetailPage.tsx", "notes": ""},
        # HR pages
        {"priority": "P1", "item": "PayrollPage.tsx — mockPayrollCycles & mockEmployeeSalaries", "owner": "Frontend", "file_ref": "frontend/src/pages/hr/PayrollPage.tsx", "notes": ""},
        {"priority": "P1", "item": "AttendancePage.tsx — mockAttendance array", "owner": "Frontend", "file_ref": "frontend/src/pages/hr/AttendancePage.tsx", "notes": ""},
        {"priority": "P1", "item": "HRDashboard.tsx — mock HR data", "owner": "Frontend", "file_ref": "frontend/src/pages/hr/HRDashboard.tsx", "notes": ""},
        {"priority": "P1", "item": "LeaveManagementPage.tsx — mock leave data", "owner": "Frontend", "file_ref": "frontend/src/pages/hr/LeaveManagementPage.tsx", "notes": ""},
        # Other pages
        {"priority": "P2", "item": "VehicleOwnerDetailPage.tsx — mock data references", "owner": "Frontend", "file_ref": "frontend/src/pages/admin/VehicleOwnerDetailPage.tsx", "notes": ""},
        {"priority": "P2", "item": "VehicleDetailPage.tsx — mock data references", "owner": "Frontend", "file_ref": "frontend/src/pages/fleet/VehicleDetailPage.tsx", "notes": ""},
        {"priority": "P2", "item": "FleetReportPage.tsx — mock report data", "owner": "Frontend", "file_ref": "frontend/src/pages/reports/FleetReportPage.tsx", "notes": ""},
        {"priority": "P2", "item": "RevenueReportPage.tsx — mock revenue data", "owner": "Frontend", "file_ref": "frontend/src/pages/reports/RevenueReportPage.tsx", "notes": ""},
        {"priority": "P2", "item": "FleetPartnersDashboard.tsx — mock partners data", "owner": "Frontend", "file_ref": "frontend/src/pages/vehicle-vendor/FleetPartnersDashboard.tsx", "notes": ""},
    ])

    doc.add_paragraph()

    add_subsection(doc, "4.2  Build & Configuration")
    create_checklist_table(doc, [
        {"priority": "P0", "item": "Frontend builds without errors (npm run build --workspace=frontend)", "owner": "Frontend", "file_ref": "frontend/vite.config.ts", "notes": ""},
        {"priority": "P0", "item": "VITE_API_URL points to production backend", "owner": "Frontend", "file_ref": "frontend/.env.production", "notes": ""},
        {"priority": "P1", "item": "VITE_GOOGLE_MAPS_API_KEY set for dispatch maps", "owner": "Frontend", "file_ref": ".env → VITE_GOOGLE_MAPS_API_KEY", "notes": ""},
        {"priority": "P1", "item": "Bundle size < 5MB (check with vite build --report)", "owner": "Frontend", "file_ref": "frontend/dist/", "notes": ""},
        {"priority": "P1", "item": "Error boundaries wrap all route-level pages", "owner": "Frontend", "file_ref": "frontend/src/", "notes": ""},
        {"priority": "P2", "item": "Source maps disabled in production build", "owner": "Frontend", "file_ref": "frontend/vite.config.ts", "notes": ""},
        {"priority": "P2", "item": "Console.log statements removed from production", "owner": "Frontend", "file_ref": "vite.config.ts → drop_console", "notes": ""},
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 5: MOBILE
    # ═══════════════════════════════════════════════════════════
    add_section_header(doc, "MOBILE APPLICATION", "5")

    create_checklist_table(doc, [
        {"priority": "P1", "item": "Expo project ID configured for push notifications", "owner": "Mobile", "file_ref": "rent-a-car-mobile/app.json", "notes": ""},
        {"priority": "P1", "item": "Company ID externalized (not hardcoded)", "owner": "Mobile", "file_ref": "rent-a-car-mobile/src/config/", "notes": "KNOWN ISSUE"},
        {"priority": "P1", "item": "API base URL points to production for release builds", "owner": "Mobile", "file_ref": "rent-a-car-mobile/src/config/api.ts", "notes": ""},
        {"priority": "P1", "item": "Deep linking scheme registered (vesla://)", "owner": "Mobile", "file_ref": "rent-a-car-mobile/app.json → scheme", "notes": ""},
        {"priority": "P1", "item": "iOS App Store submission prepared (certificates, provisioning)", "owner": "Mobile", "file_ref": "Apple Developer Console", "notes": ""},
        {"priority": "P1", "item": "Android Play Store submission prepared (signing key)", "owner": "Mobile", "file_ref": "Google Play Console", "notes": ""},
        {"priority": "P2", "item": "App icons and splash screen finalized", "owner": "Design", "file_ref": "rent-a-car-mobile/assets/", "notes": ""},
        {"priority": "P2", "item": "Offline mode tested (graceful degradation)", "owner": "Mobile", "file_ref": "—", "notes": ""},
        {"priority": "P2", "item": "Firebase real-time tracking verified on mobile", "owner": "Mobile", "file_ref": "Firebase project config", "notes": ""},
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 6: TESTING
    # ═══════════════════════════════════════════════════════════
    add_section_header(doc, "TESTING & QUALITY", "6")

    add_subsection(doc, "6.1  Automated Tests")
    create_checklist_table(doc, [
        {"priority": "P0", "item": "Unit tests pass (npm test in backend)", "owner": "QA", "file_ref": "backend/src/__tests__/", "notes": ""},
        {"priority": "P0", "item": "Security pentest suite passes (45/45)", "owner": "Security", "file_ref": "npm run test:pentest", "notes": "P0-1, P0-2, P1-3"},
        {"priority": "P0", "item": "TypeScript compiles with zero errors (npm run build:typecheck)", "owner": "Backend", "file_ref": "backend/tsconfig.json", "notes": ""},
        {"priority": "P1", "item": "Pack/role management tests pass", "owner": "QA", "file_ref": "npm run test:pack-role", "notes": ""},
        {"priority": "P1", "item": "Frontend tests pass", "owner": "QA", "file_ref": ".github/workflows/frontend-tests.yml", "notes": ""},
        {"priority": "P1", "item": "Integration tests pass", "owner": "QA", "file_ref": ".github/workflows/integration-tests.yml", "notes": ""},
        {"priority": "P1", "item": "PR quality gate passes on main branch", "owner": "QA", "file_ref": ".github/workflows/pr-quality-gate.yml", "notes": ""},
    ])

    doc.add_paragraph()

    add_subsection(doc, "6.2  Manual & Load Testing")
    create_checklist_table(doc, [
        {"priority": "P1", "item": "Critical path walkthrough: Login → Dashboard → Create Booking → Contract → Invoice → Payment", "owner": "QA", "file_ref": "—", "notes": "Manual"},
        {"priority": "P1", "item": "Multi-tenant isolation tested (Company A cannot see Company B data)", "owner": "QA", "file_ref": "npm run db:verify-isolation", "notes": ""},
        {"priority": "P1", "item": "Role-based access tested (Admin, Manager, Staff, Customer)", "owner": "QA", "file_ref": "—", "notes": "Manual"},
        {"priority": "P2", "item": "Load test: 100 concurrent users on booking flow", "owner": "QA", "file_ref": "Create: tests/load/", "notes": "k6 or Artillery"},
        {"priority": "P2", "item": "Load test: 275 model schema — query performance baseline", "owner": "QA", "file_ref": "backend/performance-baseline.json", "notes": ""},
        {"priority": "P2", "item": "Mobile app tested on iOS 16+ and Android 12+", "owner": "QA", "file_ref": "—", "notes": ""},
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 7: UAE COMPLIANCE
    # ═══════════════════════════════════════════════════════════
    add_section_header(doc, "UAE COMPLIANCE & REGULATORY", "7")

    create_checklist_table(doc, [
        {"priority": "P0", "item": "VAT 5% configuration verified and applied to all invoices", "owner": "Finance", "file_ref": "backend/src/services/compliance/", "notes": "UAE FTA"},
        {"priority": "P0", "item": "Tax Registration Number (TRN) configured per company", "owner": "Finance", "file_ref": "Company profile settings", "notes": ""},
        {"priority": "P0", "item": "E-invoicing compliant with FTA Phase 2 requirements", "owner": "Backend", "file_ref": "backend/src/middleware/einvoice-*.middleware.ts", "notes": ""},
        {"priority": "P0", "item": "Invoice format includes: TRN, VAT amount, company details", "owner": "Backend", "file_ref": "backend/src/services/contract-print.service.ts", "notes": ""},
        {"priority": "P1", "item": "Corporate Tax (CT) module configured (9% threshold)", "owner": "Finance", "file_ref": "backend/src/services/corporate-tax/", "notes": ""},
        {"priority": "P1", "item": "FTA approval workflows seeded and tested", "owner": "Backend", "file_ref": "npm run seed:fta-workflows", "notes": ""},
        {"priority": "P1", "item": "VAT return generation tested", "owner": "Finance", "file_ref": "frontend/src/pages/finance/VatReturnsPage.tsx", "notes": ""},
        {"priority": "P1", "item": "Tax calendar alerts configured", "owner": "Finance", "file_ref": "frontend/src/pages/finance/TaxCalendarPage.tsx", "notes": ""},
        {"priority": "P1", "item": "UAE Pass integration reviewed (production credentials)", "owner": "Backend", "file_ref": "docs/UAEPASS_INFRASTRUCTURE_REVIEW.md", "notes": ""},
        {"priority": "P2", "item": "Data residency: Neon DB region is ME or EU (acceptable for UAE)", "owner": "DevOps", "file_ref": "Neon Dashboard → Region", "notes": ""},
        {"priority": "P2", "item": "WPS (Wage Protection System) integration ready", "owner": "HR", "file_ref": "backend/src/middleware/wps-permissions.middleware.ts", "notes": ""},
        {"priority": "P2", "item": "Arabic language support / RTL layout verified", "owner": "Frontend", "file_ref": "frontend/src/i18n/", "notes": ""},
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 8: MONITORING & OPS
    # ═══════════════════════════════════════════════════════════
    add_section_header(doc, "MONITORING & OPERATIONS", "8")

    add_subsection(doc, "8.1  Health & Observability")
    create_checklist_table(doc, [
        {"priority": "P0", "item": "Health check endpoint returns 200 for load balancer (/api/health)", "owner": "Backend", "file_ref": "DEPLOYMENT.md → Health Endpoints", "notes": ""},
        {"priority": "P1", "item": "Detailed health endpoint auth-gated (/api/health/detailed)", "owner": "Backend", "file_ref": "backend/src/routes/ (health routes)", "notes": ""},
        {"priority": "P1", "item": "Database health check (/api/health/database)", "owner": "Backend", "file_ref": "backend/src/routes/ (health routes)", "notes": ""},
        {"priority": "P1", "item": "Error tracking service configured (Sentry DSN or equivalent)", "owner": "DevOps", "file_ref": "SENTRY_DSN env var", "notes": ""},
        {"priority": "P1", "item": "Uptime monitoring configured (UptimeRobot / Pingdom)", "owner": "DevOps", "file_ref": "External service", "notes": ""},
        {"priority": "P2", "item": "APM tool evaluated (Datadog / New Relic)", "owner": "DevOps", "file_ref": "—", "notes": ""},
    ])

    doc.add_paragraph()

    add_subsection(doc, "8.2  Backup & Disaster Recovery")
    create_checklist_table(doc, [
        {"priority": "P0", "item": "Database backup strategy documented and tested", "owner": "DevOps", "file_ref": "npm run snapshot:create", "notes": ""},
        {"priority": "P0", "item": "Backup encryption key set (BACKUP_ENCRYPTION_KEY)", "owner": "DevOps", "file_ref": "Render env", "notes": ""},
        {"priority": "P1", "item": "Daily automated snapshots configured", "owner": "DevOps", "file_ref": "npm run snapshot:daily", "notes": ""},
        {"priority": "P1", "item": "Snapshot cleanup policy (retain 30 days)", "owner": "DevOps", "file_ref": "npm run snapshot:cleanup", "notes": ""},
        {"priority": "P1", "item": "Restore tested on staging environment", "owner": "DevOps", "file_ref": "npm run snapshot:restore", "notes": ""},
        {"priority": "P1", "item": "Rollback procedure documented and rehearsed", "owner": "DevOps", "file_ref": "DEPLOYMENT.md → Rollback Procedures", "notes": ""},
    ])

    doc.add_paragraph()

    add_subsection(doc, "8.3  On-Call & Incident Response")
    create_checklist_table(doc, [
        {"priority": "P1", "item": "On-call rotation defined (primary + secondary)", "owner": "Ops", "file_ref": "Create: docs/on-call-schedule.md", "notes": ""},
        {"priority": "P1", "item": "Incident severity levels defined (S1-S4)", "owner": "Ops", "file_ref": "Create: docs/incident-response.md", "notes": ""},
        {"priority": "P1", "item": "Alert notification channels (email, Slack/WhatsApp)", "owner": "Ops", "file_ref": "—", "notes": ""},
        {"priority": "P2", "item": "Runbook for common issues (DB connection, JWT errors, CORS)", "owner": "Ops", "file_ref": "DEPLOYMENT.md → Troubleshooting", "notes": ""},
        {"priority": "P2", "item": "Post-mortem template created", "owner": "Ops", "file_ref": "Create: docs/postmortem-template.md", "notes": ""},
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 9: GO-LIVE
    # ═══════════════════════════════════════════════════════════
    add_section_header(doc, "GO-LIVE & LAUNCH", "9")

    add_subsection(doc, "9.1  Data Migration")
    create_checklist_table(doc, [
        {"priority": "P0", "item": "Data migration plan from existing systems documented", "owner": "PM", "file_ref": "Create: docs/data-migration-plan.md", "notes": ""},
        {"priority": "P0", "item": "Customer data import tested on staging", "owner": "Backend", "file_ref": "npm run db:import", "notes": ""},
        {"priority": "P1", "item": "Vehicle fleet data imported and verified", "owner": "Backend", "file_ref": "npm run db:import", "notes": ""},
        {"priority": "P1", "item": "Historical contract data migrated (if applicable)", "owner": "Backend", "file_ref": "npm run db:export / db:import", "notes": ""},
        {"priority": "P1", "item": "Data validation report generated post-import", "owner": "QA", "file_ref": "—", "notes": ""},
    ])

    doc.add_paragraph()

    add_subsection(doc, "9.2  User Acceptance & Training")
    create_checklist_table(doc, [
        {"priority": "P0", "item": "UAT completed by business stakeholders", "owner": "PM", "file_ref": "Create: docs/uat-signoff.md", "notes": ""},
        {"priority": "P0", "item": "UAT sign-off document signed", "owner": "PM", "file_ref": "—", "notes": ""},
        {"priority": "P1", "item": "Admin user training conducted", "owner": "PM", "file_ref": "Create: docs/training/", "notes": ""},
        {"priority": "P1", "item": "End-user training materials prepared", "owner": "PM", "file_ref": "—", "notes": ""},
        {"priority": "P2", "item": "Video walkthroughs recorded for key workflows", "owner": "PM", "file_ref": "—", "notes": ""},
    ])

    doc.add_paragraph()

    add_subsection(doc, "9.3  Launch Sequence")
    create_checklist_table(doc, [
        {"priority": "P0", "item": "Go/No-Go meeting held — all P0 items green", "owner": "PM", "file_ref": "This document", "notes": ""},
        {"priority": "P0", "item": "Production database provisioned and seeded", "owner": "DevOps", "file_ref": "npm run seed:production", "notes": ""},
        {"priority": "P0", "item": "DNS cutover scheduled during low-traffic window", "owner": "DevOps", "file_ref": "—", "notes": ""},
        {"priority": "P1", "item": "Feature flags configured for gradual rollout", "owner": "Backend", "file_ref": "frontend/src/pages/admin/FeatureManagementPage.tsx", "notes": ""},
        {"priority": "P1", "item": "Support escalation path defined (L1 → L2 → L3)", "owner": "Ops", "file_ref": "Create: docs/escalation-matrix.md", "notes": ""},
        {"priority": "P1", "item": "Communication plan: notify customers of launch", "owner": "PM", "file_ref": "—", "notes": ""},
        {"priority": "P1", "item": "War room / rapid response team for first 48 hours", "owner": "PM", "file_ref": "—", "notes": ""},
        {"priority": "P2", "item": "Performance monitoring dashboard set up for launch day", "owner": "DevOps", "file_ref": "—", "notes": ""},
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 10: SUMMARY & SIGN-OFF
    # ═══════════════════════════════════════════════════════════
    add_section_header(doc, "SUMMARY & SIGN-OFF", "10")

    add_subsection(doc, "10.1  Priority Summary")

    summary_table = doc.add_table(rows=4, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.style = 'Table Grid'

    summary_headers = ["Priority", "Total Items", "Completed", "Remaining"]
    for i, h in enumerate(summary_headers):
        set_cell_bg(summary_table.rows[0].cells[i], HEADER_BG)
        add_cell_text(summary_table.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=11,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    summary_data = [
        ("P0 — BLOCKERS", "~25", "___", "___"),
        ("P1 — HIGH", "~60", "___", "___"),
        ("P2 — MEDIUM", "~30", "___", "___"),
    ]
    for i, (pri_label, total, done, remain) in enumerate(summary_data):
        row = summary_table.rows[i + 1]
        pri = pri_label[:2]
        pri_color, pri_bg = priority_tag(pri)
        set_cell_bg(row.cells[0], pri_bg)
        add_cell_text(row.cells[0], pri_label, bold=True, color=pri_color, size=10)
        add_cell_text(row.cells[1], total, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(row.cells[2], done, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(row.cells[3], remain, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    doc.add_paragraph()

    add_subsection(doc, "10.2  Go/No-Go Decision")

    p = doc.add_paragraph()
    run = p.add_run("Deployment is authorized when ALL P0 items are ☑ complete.")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()

    decision_table = doc.add_table(rows=2, cols=2)
    decision_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    decision_table.style = 'Table Grid'

    set_cell_bg(decision_table.rows[0].cells[0], DONE_BG)
    add_cell_text(decision_table.rows[0].cells[0], "☐  GO — All P0 resolved, launch approved", bold=True, size=12, color=RGBColor(0x2E, 0x7D, 0x32))
    set_cell_bg(decision_table.rows[0].cells[1], P0_BG)
    add_cell_text(decision_table.rows[0].cells[1], "☐  NO-GO — P0 items remain, defer launch", bold=True, size=12, color=P0_COLOR)

    set_cell_bg(decision_table.rows[1].cells[0], "F5F5F5")
    add_cell_text(decision_table.rows[1].cells[0], "Decision Date: _______________", size=10)
    set_cell_bg(decision_table.rows[1].cells[1], "F5F5F5")
    add_cell_text(decision_table.rows[1].cells[1], "Decision By: _______________", size=10)

    doc.add_paragraph()
    doc.add_paragraph()

    add_subsection(doc, "10.3  Sign-Off")

    signoff_table = doc.add_table(rows=6, cols=4)
    signoff_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    signoff_table.style = 'Table Grid'

    signoff_headers = ["Role", "Name", "Signature", "Date"]
    for i, h in enumerate(signoff_headers):
        set_cell_bg(signoff_table.rows[0].cells[i], HEADER_BG)
        add_cell_text(signoff_table.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    roles = ["Project Manager", "Technical Lead", "QA Lead", "Security Lead", "Business Owner"]
    for i, role in enumerate(roles):
        row = signoff_table.rows[i + 1]
        add_cell_text(row.cells[0], role, bold=True, size=10)
        add_cell_text(row.cells[1], "", size=10)
        add_cell_text(row.cells[2], "", size=10)
        add_cell_text(row.cells[3], "", size=10)
        # Set row height for signature space
        row.height = Cm(1.2)

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer note
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— End of Deployment Readiness Checklist —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')} | Vesla ERP v2.0")
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    return doc


if __name__ == "__main__":
    doc = build_document()
    output_path = "/Users/miguelitodeguzman/Projects/tech-project/docs/Vesla_ERP_Deployment_Readiness_Checklist.docx"
    doc.save(output_path)
    print(f"✅ Deployment Readiness Checklist saved to: {output_path}")
